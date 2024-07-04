import json
import os
import datetime
import io
import re
import win32api                                         # Read the Windows login.
import win32net                                         # Read the Windows login.
import sys
import ctypes
import time
import pythoncom
import win32com.client as win
import socket
import re as regex
import yaml
import hashlib
import shutil
import subprocess
import pyautogui                                        # Press keyboard outside the browser.
import pyscreenshot as pyscreenshot
from deep_translator import GoogleTranslator
from docx import Document
from docx.shared import Inches                          # Used to insert image in .docx file.
from docx.shared import RGBColor
from PIL import ImageGrab
from collections import Counter                         # Used in automatizationCore_GitLab.

verbs = None
logs = None


# Colored the text.
class Textcolor:
    BLUE = '\033[1;34;47m'  # Blue
    HIGHLIGHT = '\033[7m'   # Highlight
    GREEN = '\033[32m'      # Green (seems yellow)
    WARNING = '\033[93m'    # Yellow
    FAIL = '\033[91m'       # Red
    END = '\033[00m'        # End of line
    BOLD = '\033[1m'        # Bold
    UNDERLINE = '\033[4m'   # Underline


class Main:

    def __init__(self):
        pass

    # Called to clean the file 'Tokens.txt'
    def clean_token_file(self):
        name = os.getlogin()
        file_path = os.path.join(directories["TokensFile"], 'Tokens.txt')
        file = open(file_path, 'w')
        file.close()

    # Get the Windows full name.
    def get_display_name(self):
        get_user_name_ex = ctypes.windll.secur32.GetUserNameExW
        name_display = 3

        size = ctypes.pointer(ctypes.c_ulong(0))
        get_user_name_ex(name_display, None, size)

        name_buffer = ctypes.create_unicode_buffer(size.contents.value)
        get_user_name_ex(name_display, name_buffer, size)

        return name_buffer.value

    # Set language.
    def setLanguage(**kwargs):
        try:
            # kwargs variables.
            language = kwargs.get('language')

            if language == 'pt_BR':
                Main.loadConfigs(language='pt')  # Change to the Portuguese language.
                print(f"{Textcolor.GREEN}{otherConfigs['NoTranslating']}{Textcolor.END}\n")
            elif language == 'en_US':
                need_translation, new_hash = Main.configureLanguage(language='en')
                if need_translation:
                    Main.translateMsg(language='en', new_hash=new_hash)  # Translation.
                else:
                    print(f"{Textcolor.GREEN}{otherConfigs['NoTranslating']}{Textcolor.END}\n")
                Main.loadConfigs(language='en')  # Change to the English language.
            else:  # es
                need_translation, new_hash = Main.configureLanguage(language='es')
                if need_translation:
                    Main.translateMsg(language='es', new_hash=new_hash)  # Translation.
                else:
                    print(f"{Textcolor.GREEN}{otherConfigs['NoTranslating']}{Textcolor.END}\n")
                Main.loadConfigs(language='es')  # Change to the Spanish language.

            # Variables.
            otherConfigs['Language'] = language

            Main.addLogs(message="General", value=logs["SetLanguage"])

        except Exception as ex:
            print(f"{Textcolor.FAIL}{logs['ErrorSetLanguage']['Msg']}{Textcolor.END}", ex)
            Main.addLogs(message="setLanguage", value=logs["ErrorSetLanguage"]['Msg'], value1=ex)

    # Ask and save the Token in a file.
    def saveToken(**kwargs):
        try:
            # kwargs variables.
            file_path = kwargs.get("file_path")
            name = kwargs.get("name")

            TokenField.show_token_input(TokenField)
            TokenField.invalid_token_msg(TokenField)
            TokenField().run()
            otherConfigs['Token'] = TokenField().token_input_callback()

            print(otherConfigs['InvalidTokenMessage']['Msg'])
            otherConfigs['Token'] = input(otherConfigs['InformTokenPart1']['Msg'] + ' ' +
                                          otherConfigs['InformTokenPart2']['Msg'] + ' ' +
                                          directories['TokenExpiredUrl'] + ': ')

            token_file = open(file_path, 'a')
            token_file.write(str(name) + ',' + otherConfigs['Token'] + ',\n')
            token_file.close()
            Main.addLogs(message="General", value=logs["SaveToken"])

        except Exception as ex:
            print(f"{Textcolor.FAIL}{logs['ErrorSaveToken']['Msg']}{Textcolor.END}", ex)
            Main.addLogs(message="General", value=logs["ErrorSaveToken"], value1=ex)

    # Create the directories.
    def createDirectory(**kwargs):

        try:
            # kwargs variables.
            path_folder = kwargs.get("path_folder")

            create = False

            if not os.path.exists(path_folder):
                os.makedirs(path_folder)
                create = True
            else:
                # Clear the old logs and evidences (Older than 30 days).
                current_time = time.time()
                for item in os.listdir(path_folder):

                    # Delete folders or files.
                    if os.path.isdir(os.path.join(path_folder, item)):
                        creation_time = os.path.getmtime(os.path.join(path_folder, item))
                        if (current_time - creation_time) // (24 * 3600) >= 30:
                            shutil.rmtree(os.path.join(path_folder, item), ignore_errors=True)
                            Main.addLogs(message="General", value=logs["DeleteFolder"],
                                         value1=path_folder, value2=item)
                    else:
                        creation_time = os.path.getmtime(os.path.join(path_folder, item))
                        if (current_time - creation_time) // (24 * 3600) >= 30:
                            os.remove(os.path.join(path_folder, item))
                            Main.addLogs(message="General", value=logs["DeleteFile"],
                                         value1=os.path.join(path_folder, item))

            return create

        except Exception as ex:
            print(f"{Textcolor.FAIL}{logs['ErrorCreateDirectory']['Msg']}{Textcolor.END}", ex)
            Main.addLogs(message="General", value=logs["ErrorCreateDirectory"], value1=ex)

            return create

    # Delete the directories.
    def deleteDirectory(self, **kwargs):
        try:
            # kwargs variables.
            path_folder = kwargs.get('directory')

            if os.path.exists(path_folder):
                shutil.rmtree(path_folder)

        except Exception as ex:
            print(f"{Textcolor.FAIL}{logs['ErrorDeleteDirectory']['Msg']}{Textcolor.END}", ex)
            Main.addLogs(message="General", value=logs["ErrorDeleteDirectory"], value1=ex)

    # Add the screenshots in the Word file.
    def wordAddSteps(**kwargs):

        try:
            # kwargs arguments.
            test_case_id = kwargs.get('test_case_id')
            name_testcase = kwargs.get('name_testcase')
            word_path = kwargs.get('word_path')
            steps_list = kwargs['steps_list']
            test_set_path = kwargs.get('test_set_path')
            step_failed = kwargs.get('step_failed')
            executed_by = kwargs.get('executed_by')
            take_picture_status = kwargs.get('take_picture_status', True)
            completed_date = kwargs.get('completed_date')
            duration = kwargs.get('duration')

            # Variables.
            tag_paragraf = [
                {'pt_BR': 'Evidências dos passos', 'en_US': 'Evidence of the steps', 'es': 'Evidencia de los pasos'}
            ]
            image_path = ""

            # Open the document.
            document = Document(word_path)
            # Search the correct paragraph.
            paragraf = Main.wordSeachText(document=document, text=tag_paragraf[0][otherConfigs['Language']])
            # Set the variable.
            image_resize = True

            if paragraf is None:
                print('\033[31m' + '\n' + logs['ErrorWordFindParagraph']['Msg'] + '\033[0;0m')
                Main.addLogs(message="General", value=logs["ErrorWordFindParagraph"])
                return None

            # Add the info in the file.
            if not Main.wordAddInfo(document=document, test_case_id=test_case_id,
                                    name_testcase=name_testcase, step_number=len(steps_list),
                                    completed_date=completed_date, executed_by=executed_by, duration=duration):
                print('\033[31m' + '\n' + logs['ErrorWordSetCTInfo']['Msg'] + '\033[0;0m')
                Main.addLogs(message="General", value=logs["ErrorWordSetCTInfo"])

                return None

            step_order = 1

            # Read the test case steps and add them in the order.
            for step in steps_list:

                verb = step.split()[0]
                # Don't execute the step with No / Não.
                if verb not in ('"No"', '"Não"', '"No"'.replace('"', ''),
                                '"Não"'.replace('"', '')):

                    # Last step or take_picture_status is true.
                    if verb not in ('Fechar', 'Cerrar', 'Close') and take_picture_status:
                        # Check the image size.
                        image_path = os.path.join(test_set_path, otherConfigs["EvidenceName"] +
                                                  str(step_order).zfill(2) + otherConfigs["EvidenceExtension"])
                        image = ImageGrab.Image.open(image_path)

                        if image.size[0] <= 1500:
                            image_resize = False

                        step = Main.ReplacePasswordEvidence(step=step)

                    paragraf = document.add_paragraph(otherConfigs["StepName"] + " " + str(step_order) + " - " + step)

                    run_paragraf = paragraf.add_run()

                    if step_failed == step_order:
                        # Add the error message in the document.
                        paragraf = document.add_paragraph(otherConfigs["StepNotFound"]['Msg'])
                        run_paragraf = paragraf.runs[0]
                        run_paragraf.font.color.rgb = RGBColor(*(255, 0, 0))

                    # Add the comment to the Manual Evidence.
                    # if (comment is not None) and (step_failed == step_order):
                    #     # Add the error message in the document.
                    #     paragraf = document.add_paragraph(comment)
                    #     run_paragraf = paragraf.add_run()

                    if (verb not in ('Fechar', 'Cerrar', 'Close') and take_picture_status and
                            otherConfigs['APIStep'] is False):
                        # Resize the image if it is not full screen.
                        run_paragraf.add_break()
                        if image_resize:
                            # Change the print size.
                            run_paragraf.add_picture(image_path, width=eval(otherConfigs["EvidenceWidth"]),
                                                     height=eval(otherConfigs["EvidenceHeight"]))
                        else:
                            run_paragraf.add_picture(image_path, width=Inches(5))
                    else:
                        if "AUTHORIZATION" in step.upper():
                            api_evidence_step = otherConfigs['API_Authorization']
                        elif "HEADERS" in step.upper():
                            api_evidence_step = otherConfigs['API_Headers']
                        elif "BODY" in step.upper():
                            api_evidence_step = otherConfigs['API_Body']
                        elif "ENDPOINT" in step.upper():
                            api_evidence_step = otherConfigs['API_Endpoint']
                        elif "PARAMS" in step.upper():
                            api_evidence_step = otherConfigs['API_Params']
                        elif "STATUS CODE" in step.upper():
                            api_evidence_step = str(otherConfigs['StatusCodeAPI'])
                        else:  # Response.
                            api_evidence_step = json.dumps(otherConfigs['ResponseAPI'], indent=2)

                        paragraf = document.add_paragraph(api_evidence_step)
                        run_paragraf = paragraf.runs[0]
                        run_paragraf.bold = False

                else:
                    paragraf = document.add_paragraph(otherConfigs["StepName"] + " " + str(step_order) + " - " +
                                                      otherConfigs["DisabledStep"]['Msg'])
                    run_paragraf = paragraf.add_run()

                step_order += 1

            # Save the file.
            if step_failed:
                path = os.path.join(test_set_path, '[BUG] - ' +
                                    otherConfigs["ETSName"] + str(test_case_id) + " - " + str(name_testcase)
                                    + otherConfigs["ETSExtension"])
            else:
                path = os.path.join(test_set_path, otherConfigs["ETSName"] + str(test_case_id) + " - " +
                                    str(name_testcase) + otherConfigs["ETSExtension"])
            document.save(path)

        except Exception as ex:
            print(f"{Textcolor.FAIL}{logs['ErrorWordAddSteps']['Msg']}{Textcolor.END}", ex)
            Main.addLogs(message="General", value=logs["ErrorWordAddSteps"], value1=ex)
            path = None

        return path

    # Replace the password in the evidence file.
    def ReplacePasswordEvidence(**kwargs):

        try:
            # kwargs variables.
            step = kwargs.get("step")

            matchers = ['senha', 'contraseña', 'password']

            for match in matchers:
                if match in step.lower():
                    pos_password = step.lower().find(match)
                    next_space = pos_password + len(match) + 1
                    password_string = step[next_space + 1: -1]
                    if password_string == '':
                        print(f"{Textcolor.FAIL}{logs['ErrorReplacePasswordPosition']['Msg']}"
                              f"{Textcolor.END}")
                        Main.addLogs(message="General", value=logs["ErrorReplacePasswordPosition"])
                        ####exit(1)
                    step = step.replace(password_string, '*******')

            return step

        except Exception as ex:
            print(f"{Textcolor.FAIL}{logs['ErrorReplacePasswordEvidence']['Msg']}{Textcolor.END}", ex)
            Main.addLogs(message="General", value=logs["ErrorReplacePasswordEvidence"], value1=ex)

    # Search the text in the file.
    def wordSeachText(**kwargs):

        try:
            # kwargs variables.
            document = kwargs.get("document")
            text = kwargs.get("text")

            for p in document.paragraphs:
                if p.text == text:
                    return p
        except Exception as ex:
            print(f"{Textcolor.FAIL}{logs['ErrorWordSeachText']['Msg']}{Textcolor.END}", ex)
            Main.addLogs(message="General", value=logs["ErrorWordSeachText"], value1=ex)

        return None

    # Add the test case info in the Word file.
    def wordAddInfo(**kwargs):

        # kwargs arguments.
        document = kwargs.get('document')
        test_case_id = kwargs.get('test_case_id')
        name_testcase = kwargs.get('name_testcase')
        step_number = kwargs.get('step_number')
        executed_by = kwargs.get('executed_by')
        completed_date = kwargs.get('completed_date')
        duration = kwargs.get('duration')

        try:

            tag_language = [
                {'pt_BR': 'ID do Caso de Teste: ', 'en_US': 'Test Case ID: ', 'es': 'ID del Prueba: '},
                {'pt_BR': 'Caso de Teste: ', 'en_US': 'Test Case: ', 'es': 'Prueba: '},
                {'pt_BR': 'Teste executado por: ', 'en_US': 'Test executed by: ', 'es': 'Prueba realizada por: '},
                {'pt_BR': 'Evidência gerada por: ', 'en_US': 'Evidence generated by: ', 'es': 'Evidencia generada por: '},
                {'pt_BR': 'Data de Execução: ', 'en_US': 'Execution date: ', 'es': 'Fecha de ejecución: '},
                {'pt_BR': 'Total de passos: ', 'en_US': 'Total steps: ', 'es': 'Pasos totales: '},
                {'pt_BR': 'Duração da execução do teste: ', 'en_US': 'Test execution time: ', 'es': 'Tiempo de ejecución de la prueba: '}
            ]

            # Test case GitLab ID.
            control = Main.wordSeachText(document=document, text=tag_language[0][otherConfigs['Language']])
            control.add_run(str(test_case_id)).bold = True

            # Test case name.
            control = Main.wordSeachText(document=document, text=tag_language[1][otherConfigs['Language']])
            control.add_run(name_testcase).bold = True

            # Executed by.
            control = Main.wordSeachText(document=document, text=tag_language[2][otherConfigs['Language']])
            control.add_run(executed_by).bold = True

            # Evidence generate by.
            control = Main.wordSeachText(document=document, text=tag_language[3][otherConfigs['Language']])
            control.add_run(executed_by).bold = True

            # Execution date.
            control = Main.wordSeachText(document=document, text=tag_language[4][otherConfigs['Language']])
            control.add_run(str(completed_date)).bold = True

            # Total steps.
            control = Main.wordSeachText(document=document, text=tag_language[5][otherConfigs['Language']])
            control.add_run(str(step_number)).bold = True

            # Test execution duration.
            control = Main.wordSeachText(document=document, text=tag_language[6][otherConfigs['Language']])
            control.add_run(duration).bold = True

            return True

        except Exception as ex:
            print(f"{Textcolor.FAIL}{logs['ErrorWordAddInfo']['Msg']}{Textcolor.END}", ex)
            Main.addLogs(message="General", value=logs["ErrorWordAddInfo"], value1=ex)

            return False

    # Function to convert docx to pdf.
    def wordToPDF(**kwargs):

        try:
            # kwargs variables.
            path = kwargs.get("path")

            word_format_pdf = 17

            # Initialize.
            pythoncom.CoInitialize()

            word = win.Dispatch('Word.Application')
            document = word.Documents.Open(path)

            pdf_path = path.replace("docx", "pdf")
            document.SaveAs(pdf_path, FileFormat=word_format_pdf)
            document.Close()

            word.Quit()

            return pdf_path

        except Exception as ex:
            print(f"{Textcolor.FAIL}{logs['ErrorWordToPDF']['Msg']}{Textcolor.END}", ex)
            Main.addLogs(message="General", value=logs["ErrorWordToPDF"], value1=ex)
            return None

    # Delete files.
    def deleteFiles(**kwargs):

        try:
            # kwargs arguments.
            file_path = kwargs.get('path_log')
            extension = kwargs.get('extension')
            exact_file = kwargs.get('exact_file')

            # Delete all the files in a directory with the specific extension OR all if the extension is '*'.
            if file_path:

                files = os.listdir(file_path)

                for item in files:
                    if item.endswith(extension) or extension == '*':
                        os.remove(os.path.join(file_path, item))

            else:
                os.remove(exact_file)

        except Exception as ex:
            print(f"{Textcolor.FAIL}{logs['ErrorDeleteFiles']['Msg']}{Textcolor.END}", ex)
            Main.addLogs(message="General", value=logs["ErrorDeleteFiles"], value1=ex)

    # Add the log in the file.
    def addLogs(**kwargs):

        try:
            # kwargs variables.
            message = kwargs.get("message")
            value = kwargs.get("value")
            value1 = kwargs.get("value1")
            value2 = kwargs.get("value2")

            datetime_log: str = datetime.datetime.now().strftime("%d/%m/%y %H:%M:%S")

            # Get the hostname.
            otherConfigs["ComputerName"] = socket.gethostname()
            hostname = otherConfigs["ComputerName"]
            # Get the datetime.
            date_log = str(datetime.datetime.now().strftime("%d.%m.%Y"))
            # Set the file name.
            path = os.path.join(directories["LogFolder"], hostname + " - " + date_log + ".log")

            if not os.path.isdir(directories["LogFolder"]):
                Main.createDirectory(path_folder=directories["LogFolder"])

            # Append the log file.
            with open(path, 'a+', encoding='utf-8') as log_file:
                if message == 'NewConfig':  # Set the first line.
                    log_file.write("\n " + "*" * 61 + datetime_log + "*" * 61 + "\n")
                    log_file.write("\nLOAD INFORMATION'S FROM GITLAB\n")
                elif message == 'NewSession':
                    log_file.write(str(value) + "\n")  # Test case name.
                    log_file.write("")
                elif message == 'General':
                    type_log = value['Type']
                    message_log = value['Msg']
                    local_log = value['Where']

                    # Alignment.
                    type_log = "{:<9}".format(type_log)

                    # Change the error origin if informed.
                    if value1 is not None and value2 is None:
                        msg_format = "{:<15}"
                        log_file.write(datetime_log + " - " + type_log + " - " + msg_format.format(message_log) + " - " + value1 + "\n")
                    elif value1 is not None and value2 is not None:
                        log_file.write(datetime_log + " - " + type_log + " - '" + value1 + "' " + message_log + " '" +
                                       value2 + "' - " + local_log + "\n")
                    else:
                        log_file.write(datetime_log + " - " + type_log + " - " + message_log + " - " + local_log + "\n")

                elif message == 'EndExecution':
                    log_file.write("\n " + "*" * 138 + "\n")
                else:
                    log_file.write(datetime_log + " - Log       - " + value + " " + " - " + str(value1) + "\n")

        except Exception as ex:
            Main.addLogs(message="General", value=logs["ErrorAddlog"]['Msg'], value1=str(ex))
            print(f"{Textcolor.FAIL}{logs['ErrorAddlog']['Msg']}{Textcolor.END}", ex)

    # Remove the HTML from the string.
    def removeHTML(**kwargs):

        try:
            # kwargs variables.
            value = kwargs.get("value")

            pattern = regex.compile('<.*?>')
            value = regex.sub(pattern, '', value).strip()

            value = value.replace('&nbsp;', ' ')
            value = value.replace('&lt;', '<')
            value = value.replace('&gt;', '>')

            return value

        except Exception as ex:

            print(f"{Textcolor.FAIL}{logs['ErrorRemoveHTML']['Msg']}{Textcolor.END}", ex)
            Main.addLogs(message="General", value=logs["ErrorRemoveHTML"]['Msg'], value1=str(ex))

    # Translate the messages.
    def translateMsg(**kwargs):
        try:
            # kwargs arguments.
            language = kwargs.get('language')
            new_hash = kwargs.get('new_hash')

            # Paths
            path_origin_yml = os.path.join(os.getcwd(), 'Automation', 'configs', 'dictionary-pt.yml')
            path_translated_yml = os.path.join(os.getcwd(), 'Automation', 'configs', 'dictionary-' + language + '.yml')

            regex_pattern = regex.compile('(.*?,)(.*?,)(.*)')

            tag = 'Msg:'

            if os.path.isfile(path_translated_yml):
                Main.deleteFiles(exact_file=path_translated_yml)

            with open(path_origin_yml, 'r') as yml_file:
                lines = yml_file.readlines()
                for line in lines:
                    if tag in line:
                        with open(path_translated_yml, 'a', encoding='utf-8') as yml_new_file:
                            group_type = regex.match(regex_pattern, line).group(1)
                            group_msg = regex.match(regex_pattern, line).group(2)
                            msg_translated = GoogleTranslator(source='pt', target=language).\
                                translate((group_msg[len(tag) + 1:len(group_msg) - 1]).strip())
                            group_where = regex.match(regex_pattern, line).group(3)
                            yml_new_file.write(f'{group_type}{tag} {msg_translated},{group_where}\n')
                    else:
                        with open(path_translated_yml, 'a') as yml_new_file:
                            yml_new_file.write(f'{line}')

            Main.saveHash(new_hash=new_hash, path_part=language)
            print(f"{Textcolor.GREEN}{otherConfigs['TranslateMessage']}{Textcolor.END}")

        except Exception as ex:
            print(f"{Textcolor.FAIL}{logs['ErrorTranslateMessage']['Msg']}{Textcolor.END}", ex)
            Main.addLogs(message="General", value=logs["ErrorTranslateMessage"], value1=ex)

    # Configure the language for the automation.
    def configureLanguage(**kwargs):
        try:
            need_translation = False
            language = kwargs.get('language')

            path_file = os.path.join(os.getcwd(), directories["ConfigFolder"], 'dictionary-pt.yml')
            new_hash = Main.generateHash(path_file=path_file)
            actual_hash = Main.readHash(directory=directories["HashFolder"], language=language,
                                        actual_file='hash_dictionary.txt')

            if new_hash != actual_hash:
                need_translation = True

            return need_translation, new_hash

        except Exception as ex:
            print(f"{Textcolor.FAIL}{logs['ErrorConfigureLanguage']['Msg']}{Textcolor.END}", ex)
            Main.addLogs(message="General", value=logs["ErrorConfigureLanguage"], value1=ex)

    # Generate the hash for a file.
    def generateHash(**kwargs):
        try:
            # kwargs arguments.
            path_file = kwargs.get('path_file')

            BLOCKSIZE = 65536
            hasher = hashlib.sha1()
            with open(path_file, 'rb') as target_file:
                buffer = target_file.read(BLOCKSIZE)
                while len(buffer) > 0:
                    hasher.update(buffer)
                    buffer = target_file.read(BLOCKSIZE)

            return hasher.hexdigest()

        except Exception as ex:
            print(f"{Textcolor.FAIL}{logs['ErrorGenerateHash']['Msg']}{Textcolor.END}", ex)
            Main.addLogs(message="General", value=logs["ErrorGenerateHash"], value1=ex)

    # Read the hash in a file.
    def readHash(**kwargs):
        try:
            # kwargs arguments.
            directory = kwargs.get('directory')
            language = kwargs.get('language')
            actual_file = kwargs.get('actual_file')

            file_read = os.path.join(directory, language + '-' + actual_file)

            if os.path.exists(file_read):

                with open(file_read, 'r') as file:
                    content = file.readline()
            else:
                content = None

            return content

        except Exception as ex:
            print(f"{Textcolor.FAIL}{logs['ErrorReadHash']['Msg']}{Textcolor.END}", ex)
            Main.addLogs(message="General", value=logs["ErrorReadHash"], value1=ex)

    # Save the hash in a file.
    def saveHash(**kwargs):
        try:
            # kwargs arguments.
            new_hash = kwargs.get('new_hash')
            path_part = kwargs.get('path_part')

            # Variables.
            hash_file_path = os.path.join(directories["HashFolder"], path_part + '-hash_dictionary.txt')

            if not os.path.exists(directories['HashFolder']):
                Main.createDirectory(path_folder=directories['HashFolder'])
            with open(hash_file_path, 'w') as hash_file:
                hash_file.write(new_hash)

        except Exception as ex:
            print(f"{Textcolor.FAIL}{logs['ErrorSaveHash']['Msg']}{Textcolor.END}", ex)
            Main.addLogs(message="General", value=logs["ErrorSaveHash"], value1=ex)

    # Load the configuration file.
    def loadConfigs(**kwargs):

        # kwargs parameters.
        language = kwargs.get('language')

        try:
            global verbs, logs, directories, otherConfigs, searchForAttribute, searchForComponent

            # yml path.
            path = os.path.join(os.getcwd(), 'Automation', 'configs', 'dictionary-' + language + '.yml')
            if not os.path.exists(path):  # For the exe file.
                path = os.path.join(os.getcwd(), 'configs', 'dictionary-' + language + '.yml')

            with open(path, encoding='utf-8') as configFile:
                config = yaml.safe_load(configFile)

            # Add the sections.
            verbs = config['verbs']
            logs = config['logs']
            directories = config['directories']
            otherConfigs = config['otherConfigs']
            searchForAttribute = config["searchForAttribute"]
            searchForComponent = config["searchForComponent"]

        except Exception as ex:
            print(f"{Textcolor.FAIL}{logs['ErrorLoadConfigs']['Msg']}{Textcolor.END}", ex)
            Main.addLogs(message="General", value=logs["ErrorLoadConfigs"], value1=str(ex))

    # Calculate the percentage of execution.
    def percentage(**kwargs):
        try:
            # kwargs parameters.
            actual_number = kwargs.get('actual')
            total = kwargs.get('total')

            # line = "*" * 12
            percentage = "{0:.2f}".format((actual_number / total) * 100)

            print(f"{Textcolor.BLUE}{' '}"
                  f"{otherConfigs['Percentage']['Msg']}"
                  f"{' '}{Textcolor.UNDERLINE}{percentage}{'%'}{' '}{Textcolor.END}"
                  f"{Textcolor.BLUE}{Textcolor.END}\n")

            Main.addLogs(message="General", value=logs["Percentage"])

        except Exception as ex:
            print(f"{Textcolor.FAIL}{logs['ErrorPercentage']['Msg']}{Textcolor.END}", ex)
            Main.addLogs(message="General", value=logs["ErrorPercentage"], value1=str(ex))

    # Verify if the file exist.
    def verifyFile(self, **kwargs):
        try:

            # kwargs parameters.
            path = kwargs.get('path')
            extension = kwargs.get('extension')
            msg_not_found = kwargs.get('msg_not_found')
            msg_found = kwargs.get('msg_found')

            files = os.listdir(path)
            for file in files:
                if file.endswith(extension):
                    print(f"{Textcolor.GREEN}{msg_found}{Textcolor.END}")
                    return True
                else:
                    print(f"{Textcolor.GREEN}{msg_not_found}{Textcolor.END}")
                    return False

        except Exception as ex:
            print(f"{Textcolor.FAIL}{logs['ErrorVerifyFile']['Msg']}{Textcolor.END}", ex)
            Main.addLogs(message="General", value=logs["ErrorVerifyFile"], value1=str(ex))

    # Compare files using Beyond Compare.
    def compareBeyondCompare(**kwargs):
        try:
            # kwargs parameters.
            baseline = kwargs.get('baseline')
            new_file = kwargs.get('new_file')
            test_name = kwargs.get('test_name')

            baseline = os.path.join(directories['CompareDownloadFolder'], test_name, baseline)
            new_file = os.path.join(directories['CompareDownloadFolder'], test_name, new_file)

            # Update the Settings.
            subprocess.Popen([directories['BeyondCompare'], directories['BeyondCompareSettings'], '/silent'])

            file_session = Main._checkSessionBCFile(baseline=baseline)

            # Open the session.
            subprocess.Popen([directories['BeyondCompare'], file_session, '/silent'])

            # Disable the mouse.
            pyautogui.FAILSAFE = False

            # Open the file (Left).
            time.sleep(10)
            pyautogui.hotkey('ctrl', 'o', interval=.5)
            pyautogui.typewrite(baseline)
            pyautogui.hotkey('Enter', interval=.5)

            pyautogui.press('tab', interval=.5)

            # Open the file (Right).
            pyautogui.hotkey('ctrl', 'o', interval=.5)
            pyautogui.typewrite(new_file)
            pyautogui.hotkey('Enter', interval=.5)

            # Enable the edit option.
            pyautogui.press('f2', interval=.5)

            # Disable the mouse.
            pyautogui.FAILSAFE = True

            while Main.checkProcess(process='BCompare.exe'):
                time.sleep(1)

            print(f"{Textcolor.GREEN}{test_name} - {logs['CompareFile']['Msg']}{Textcolor.END}")

        except Exception as ex:
            print(f"{Textcolor.FAIL}{logs['ErrorCompareFile']['Msg']}{Textcolor.END}", ex)
            Main.addLogs(message="General", value=logs["ErrorCompareFile"], value1=str(ex))

    # Check if the test case is a Desktop test case.
    def _checkDesktop_TC(**kwargs):

        # kwargs variable.
        list_steps = kwargs.get("list_steps")

        desktop_tc = False

        for cont, _ in enumerate(otherConfigs['DesktopFunctions']):
            if any(otherConfigs['DesktopFunctions'][cont] in desktop_verb for desktop_verb in list_steps):
                desktop_tc = True

        return desktop_tc

    @staticmethod
    # Inform the updates.
    def releaseNotes():

        path = os.path.join(os.getcwd(), 'README.md')
        release_infos = []

        if path:
            with open(path, 'r', encoding='utf-8') as readme:

                for line in readme:
                    if 'Version' in line:
                        local_version = line[10:-3]
                    if '<em>' in line:
                        date_version = line[4:-6]
                    if '</font>' in line:
                        release_infos.append(line[38:].strip())
                    if '#' in line:
                        break

        return local_version, date_version, release_infos

    # Validate content inside JSON content.
    def find_content_json(self, **kwargs):

        # kwargs variable.
        param = kwargs.get("param")
        tag = kwargs.get("tag")

        validation = False

        try:
            response = otherConfigs['ResponseAPI']

            # Tag's validation.
            for dict_id, _ in enumerate(response):
                if str(response[dict_id][tag]) == str(param):
                    validation = (validation or True)
                else:
                    validation = (validation or False)

                if validation:
                    return "Passed"
                else:
                    print(f"{Textcolor.FAIL}{logs['ErrorValidationAPI']['Msg']}{Textcolor.END}")
                    Main.addLogs(message="General", value=logs["ErrorValidationAPI"])
                    return "Failed"

        except Exception as ex:
            print(f"{Textcolor.FAIL}{logs['ErrorFindContentAPI']['Msg']}{Textcolor.END}", ex)
            Main.addLogs(message="General", value=logs["ErrorFindContentAPI"], value1=str(ex))

            return "Failed"

    def convert_seconds_to_string(self, **kwargs):

        # kwargs variable.
        time_spent = kwargs.get("time_spent")

        total_minutes = time_spent / 60.0
        minutes = int(total_minutes)
        remaining_seconds = total_minutes - (minutes * 60)

        return f"{minutes:02}:{remaining_seconds:06.3f}"
