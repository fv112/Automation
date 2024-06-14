import os
import datetime
import re
# import requests ###

import win32api                                         # Read the Windows login.
import win32net                                         # Read the Windows login.
import sys
import ctypes
import time
import pythoncom
import win32com.client as win
import socket
import re as regex
import yaml
import hashlib
import shutil
import subprocess
import pyautogui                                        # Press keyboard outside the browser.
import pyscreenshot as pyscreenshot
from requests.auth import HTTPBasicAuth
from deep_translator import GoogleTranslator
from docx import Document
from docx.shared import Inches                          # Used to insert image in .docx file.
from PIL import ImageGrab
from collections import Counter                         # Used in automatizationCore_Azure.

replaceEvidence = None
searchForAttribute = []
searchForComponent = []
messages = []


# Colored the text.
class Textcolor:
    BLUE = '\033[1;34;47m'  # Blue
    GREEN = '\033[32m'      # Green (seems yellow)
    WARNING = '\033[93m'    # Yellow
    FAIL = '\033[91m'       # Red
    END = '\033[00m'        # End of line
    BOLD = '\033[1m'        # Bold
    UNDERLINE = '\033[4m'   # Underline


class Main:

    # Called to clean the file 'Tokens.txt'
    def clean_token_file(self):
        name = os.getlogin()
        file_path = os.path.join(directories["TokensFile"], 'Tokens.txt')
        file = open(file_path, 'w')
        file.close()

    # Get the Windows full name.
    def get_display_name(self):
        get_user_name_ex = ctypes.windll.secur32.GetUserNameExW
        name_display = 3

        size = ctypes.pointer(ctypes.c_ulong(0))
        get_user_name_ex(name_display, None, size)

        name_buffer = ctypes.create_unicode_buffer(size.contents.value)
        get_user_name_ex(name_display, name_buffer, size)

        return name_buffer.value

    # Set language.
    def setLanguage(**kwargs):
        try:
            # kwargs variables.
            language = kwargs.get('language')

            if language == 'pt_BR':
                Main.loadConfigs(language='pt')  # Change to the Portuguese language.
                print(f"{Textcolor.GREEN}{otherConfigs['NoTranslating']}{Textcolor.END}\n")
            elif language == 'en_US':
                need_translation, new_hash = Main.configureLanguage(language='en')
                if need_translation:
                    Main.translateMsg(language='en', new_hash=new_hash)  # Translation.
                else:
                    print(f"{Textcolor.GREEN}{otherConfigs['NoTranslating']}{Textcolor.END}\n")
                Main.loadConfigs(language='en')  # Change to the English language.
            else:  # es
                need_translation, new_hash = Main.configureLanguage(language='es')
                if need_translation:
                    Main.translateMsg(language='es', new_hash=new_hash)  # Translation.
                else:
                    print(f"{Textcolor.GREEN}{otherConfigs['NoTranslating']}{Textcolor.END}\n")
                Main.loadConfigs(language='es')  # Change to the Spanish language.

            # Variables.
            otherConfigs['Language'] = language

            Main.addLogs(message="General", value=logs["SetLanguage"])

        except Exception as ex:
            print(f"{Textcolor.FAIL}{logs['ErrorSetLanguage']['Msg']}{Textcolor.END}", ex)
            Main.addLogs(message="setLanguage", value=logs["ErrorSetLanguage"]['Msg'], value1=ex)

    # Request the token or verify if it was informed.
    def accessAzure(self):
        try:
            name = os.getlogin()

            file_path = os.path.join(directories["TokensFile"], 'Tokens.txt')

            # Create the folder to save the token.
            if not os.path.exists(os.path.join(directories["TokensFile"])):
                os.makedirs(directories["TokensFile"])
                # Create the file and set the mode.
                with open(file_path, 'w'):
                    Main.saveToken(file_path=file_path, name=name)

            else:
                token_exist = False
                with open(file_path, 'r') as myfile:
                    for line in myfile:
                        if name in line:
                            otherConfigs['Token'] = line.split(',')[1]
                            token_exist = True
                    if token_exist is False:
                        Main.saveToken(file_path=file_path, name=name)

            otherConfigs['HttpBasicAuth'] = HTTPBasicAuth('', otherConfigs['Token'])
            Main.addLogs(message="NewConfig", value=logs["AccessAzure"]['Msg'])

        except Exception as ex:
            print(f"{Textcolor.FAIL}{logs['ErrorAccessAzure']['Msg']}{Textcolor.END}", ex)
            Main.addLogs(message="NewConfig", value=logs["ErrorAccessAzure"]['Msg'], value1=ex)

    # Ask and save the Token in a file.
    def saveToken(**kwargs):
        try:
            # kwargs variables.
            file_path = kwargs.get("file_path")
            name = kwargs.get("name")

            if otherConfigs['Interface']:
                from AppAutomation import TokenField

                TokenField.show_token_input(TokenField)
                TokenField.invalid_token_msg(TokenField)
                TokenField().run()
                otherConfigs['Token'] = TokenField().token_input_callback()

            else:
                print(otherConfigs['InvalidTokenMessage']['Msg'])
                otherConfigs['Token'] = input(otherConfigs['InformTokenPart1']['Msg'] + ' ' +
                                              otherConfigs['InformTokenPart2']['Msg'] + ' ' +
                                              directories['TokenExpiredUrl'] + ': ')

            token_file = open(file_path, 'a')
            token_file.write(str(name) + ',' + otherConfigs['Token'] + ',\n')
            token_file.close()
            Main.addLogs(message="General", value=logs["SaveToken"])

        except Exception as ex:
            print(f"{Textcolor.FAIL}{logs['ErrorSaveToken']['Msg']}{Textcolor.END}", ex)
            Main.addLogs(message="General", value=logs["ErrorSaveToken"], value1=ex)

    # Validate test case name.
    # def validateTestName(**kwargs):
    #     try:
    #         # kwargs variables.
    #         name_testcase = kwargs.get("name_testcase")
    #
    #         validation_status = None
    #
    #         validation = regex.match('.*[\.\@\!\#\$\%^\&\*\<\>\?\\\/\\|\"}{:].*', name_testcase)
    #         if validation:
    #             print(f"{Textcolor.FAIL}{logs['ErrorSpecialCharacter']['Msg']} "
    #                   f"{otherConfigs['InvalidCharacter']} {Textcolor.END}")
    #             Main.addLogs(message="General", value=logs["ErrorSpecialCharacter"],
    #                          value1=f"{otherConfigs['InvalidCharacter']}")
    #             validation_status = True
    #             ###exit(0)
    #
    #         if len(name_testcase) >= 85:
    #             print(f"{Textcolor.FAIL}{logs['ErrorSizeName']['Msg']}{Textcolor.END}")
    #             Main.addLogs(message="General", value=logs["ErrorSizeName"])
    #             validation_status = True
    #             ###exit(0)
    #
    #     except Exception as ex:
    #         ###exit(0)
    #         print(f"{Textcolor.FAIL}{logs['ErrorTestCaseValidation']['Msg']}{Textcolor.END}", ex)
    #         Main.addLogs(message="General", value=logs["ErrorTestCaseValidation"], value1=str(ex))
    #
    #     finally:
    #         return validation_status

    # Create the directories.
    def createDirectory(**kwargs):

        try:
            # kwargs variables.
            path_folder = kwargs.get("path_folder")

            create = False

            if not os.path.exists(path_folder):
                os.makedirs(path_folder)
                create = True
            else:
                # Clear the old logs and evidences (Older than 30 days).
                current_time = time.time()
                for item in os.listdir(path_folder):

                    # Delete folders or files.
                    if os.path.isdir(os.path.join(path_folder, item)):
                        creation_time = os.path.getmtime(os.path.join(path_folder, item))
                        if (current_time - creation_time) // (24 * 3600) >= 30:
                            shutil.rmtree(os.path.join(path_folder, item), ignore_errors=True)
                            Main.addLogs(message="General", value=logs["DeleteFolder"],
                                         value1=path_folder, value2=item)
                    else:
                        creation_time = os.path.getmtime(os.path.join(path_folder, item))
                        if (current_time - creation_time) // (24 * 3600) >= 30:
                            os.remove(os.path.join(path_folder, item))
                            Main.addLogs(message="General", value=logs["DeleteFile"],
                                         value1=os.path.join(path_folder, item))

        except Exception as ex:
            print(f"{Textcolor.FAIL}{logs['ErrorCreateDirectory']['Msg']}{Textcolor.END}", ex)
            Main.addLogs(message="General", value=logs["ErrorCreateDirectory"], value1=ex)

        return create

    # Delete the directories.
    def deleteDirectory(**kwargs):
        try:
            # kwargs variables.
            path_folder = kwargs.get('directory')

            if os.path.exists(path_folder):
                shutil.rmtree(path_folder)

        except Exception as ex:
            print(f"{Textcolor.FAIL}{logs['ErrorDeleteDirectory']['Msg']}{Textcolor.END}", ex)
            Main.addLogs(message="General", value=logs["ErrorDeleteDirectory"], value1=ex)

    # Add the screenshots in the Word file.
    def wordAddSteps(**kwargs):

        try:

            # kwargs arguments.
            test_case_id = kwargs.get('test_case_id')
            name_testcase = kwargs.get('name_testcase')
            word_path = kwargs.get('word_path')
            steps_list = kwargs['steps_list']
            test_set_path = kwargs.get('test_set_path')
            step_failed = kwargs.get('step_failed')
            # full_name_run_evidence = kwargs.get('full_name_run_evidence')
            take_picture_status = kwargs.get('take_picture_status', True)
            completed_date = kwargs.get('completed_date')

            # Variables.
            tag_paragraf = [
                {'pt_BR': 'Evidências dos passos', 'en_US': 'Evidence of the steps', 'es': 'Evidencia de los pasos'}
            ]
            image_path = ""

            # Open the document.
            document = Document(word_path)
            # Search the correct paragraph.
            paragraf = Main.wordSeachText(document=document, text=tag_paragraf[0][otherConfigs['Language']])
            # Set the variable.
            image_resize = True

            if paragraf is None:
                print('\033[31m' + '\n' + logs['ErrorWordFindParagraph']['Msg'] + '\033[0;0m')
                Main.addLogs(message="General", value=logs["ErrorWordFindParagraph"])
                return None

            # Add the info in the file.
            if not Main.wordAddInfo(document=document, test_case_id=test_case_id,
                                    name_testcase=name_testcase, step_number=len(steps_list),
                                    completed_date=completed_date):
                print('\033[31m' + '\n' + logs['ErrorWordSetCTInfo']['Msg'] + '\033[0;0m')
                Main.addLogs(message="General", value=logs["ErrorWordSetCTInfo"])
                return None

            step_order = 1

            # Read the test case steps and add them in the order.
            for step in steps_list:

                verb = step.split()[0]
                # Don't execute the step with No / Não.
                if verb not in ('"No"', '"Não"', '"No"'.replace('"', ''), '"Não"'.replace('"', '')):

                    # Last step or take_picture_status is true.
                    if verb not in ('Fechar', 'Cerrar', 'Close') and take_picture_status:
                        # Check the image size.
                        image_path = os.path.join(test_set_path, otherConfigs["EvidenceName"] +
                                                  str(step_order).zfill(2) + otherConfigs["EvidenceExtension"])
                        image = ImageGrab.Image.open(image_path)

                        if image.size[0] <= 1500:
                            image_resize = False

                        step = Main.ReplacePasswordEvidence(step=step)

                    paragraf = document.add_paragraph(otherConfigs["StepName"] + " " + str(step_order) + " - " + step)

                    run_paragraf = paragraf.add_run()

                    if step_failed == step_order:
                        # Add the error message in the document.
                        paragraf = document.add_paragraph(otherConfigs["StepNotFound"]['Msg'])
                        run_paragraf = paragraf.add_run()

                    # Add the comment to the Manual Evidence.
                    # if (comment is not None) and (step_failed == step_order):
                    #     # Add the error message in the document.
                    #     paragraf = document.add_paragraph(comment)
                    #     run_paragraf = paragraf.add_run()

                    if verb not in ('Fechar', 'Cerrar', 'Close') and take_picture_status:
                        # Resize the image if it is not full screen.
                        run_paragraf.add_break()
                        if image_resize:
                            # Change the print size.
                            run_paragraf.add_picture(image_path, width=eval(otherConfigs["EvidenceWidth"]),
                                                     height=eval(otherConfigs["EvidenceHeight"]))
                        else:
                            run_paragraf.add_picture(image_path, width=Inches(5))

                else:
                    paragraf = document.add_paragraph(otherConfigs["StepName"] + " " + str(step_order) + " - " +
                                                      otherConfigs["DisabledStep"]['Msg'])
                    # run_paragraf = paragraf.add_run()

                step_order += 1

            # Save the file.
            path = os.path.join(test_set_path, otherConfigs["ETSName"] + str(test_case_id) + " - " + str(name_testcase)
                                + otherConfigs["ETSExtension"])
            document.save(path)

        except Exception as ex:
            print(f"{Textcolor.FAIL}{logs['ErrorWordAddSteps']['Msg']}{Textcolor.END}", ex)
            Main.addLogs(message="General", value=logs["ErrorWordAddSteps"], value1=ex)
            path = None

        return path

    # Replace the password in the evidence file.
    def ReplacePasswordEvidence(**kwargs):

        try:
            # kwargs variables.
            step = kwargs.get("step")

            matchers = ['senha', 'contraseña', 'password']

            for match in matchers:
                if match in step.lower():
                    pos_password = step.lower().find(match)
                    next_space = pos_password + len(match) + 1
                    password_string = step[next_space + 1: -1]
                    if password_string == '':
                        print(f"{Textcolor.FAIL}{logs['ErrorReplacePasswordPosition']['Msg']}"
                              f"{Textcolor.END}")
                        Main.addLogs(message="General", value=logs["ErrorReplacePasswordPosition"])
                        ###exit(1)
                    step = step.replace(password_string, '*******')

            return step

        except Exception as ex:
            print(f"{Textcolor.FAIL}{logs['ErrorReplacePasswordEvidence']['Msg']}{Textcolor.END}", ex)
            Main.addLogs(message="General", value=logs["ErrorReplacePasswordEvidence"], value1=ex)

    # Search the text in the file.
    def wordSeachText(**kwargs):

        try:
            # kwargs variables.
            document = kwargs.get("document")
            text = kwargs.get("text")

            for p in document.paragraphs:
                if p.text == text:
                    return p
        except Exception as ex:
            print(f"{Textcolor.FAIL}{logs['ErrorWordSeachText']['Msg']}{Textcolor.END}", ex)
            Main.addLogs(message="General", value=logs["ErrorWordSeachText"], value1=ex)

        return None

    # Add the test case info in the Word file.
    def wordAddInfo(**kwargs):
        try:
            # kwargs arguments.
            document = kwargs.get('document')
            test_run_id = kwargs.get('test_run_id')
            test_case_id = kwargs.get('test_case_id')
            name_testcase = kwargs.get('name_testcase')
            summary = kwargs.get('summary')
            step_number = kwargs.get('step_number')
            full_name_run_evidence = kwargs.get('full_name_run_evidence')
            full_name_run_test = kwargs.get('full_name_run_test')
            completed_date = kwargs.get('completed_date')

            tag_language = [
                {'pt_BR': 'ID da Run: ', 'en_US': 'Run ID: ', 'es': 'ID Run: '},
                {'pt_BR': 'ID do Caso de Teste: ', 'en_US': 'Test Case ID: ', 'es': 'ID del Prueba: '},
                {'pt_BR': 'Caso de Teste: ', 'en_US': 'Test Case: ', 'es': 'Prueba: '},
                {'pt_BR': 'Descrição: ', 'en_US': 'Description: ', 'es': 'Descripción: '},
                {'pt_BR': 'Teste executado por: ', 'en_US': 'Test executed by: ', 'es': 'Prueba realizada por: '},
                {'pt_BR': 'Evidência gerada por: ', 'en_US': 'Evidence generated by: ', 'es': 'Evidencia generada por: '},
                {'pt_BR': 'Data de Execução: ', 'en_US': 'Execution date: ', 'es': 'Fecha de ejecución: '},
                {'pt_BR': 'Total de passos: ', 'en_US': 'Total steps: ', 'es': 'Pasos totales: '}
            ]

            # TestRun id.
            control = Main.wordSeachText(document=document, text=tag_language[0][otherConfigs['Language']])
            control.add_run(str(test_run_id)).bold = True

            # CT id.
            control = Main.wordSeachText(document=document, text=tag_language[1][otherConfigs['Language']])
            control.add_run(str(test_case_id)).bold = True

            # CT name.
            control = Main.wordSeachText(document=document, text=tag_language[2][otherConfigs['Language']])
            control.add_run(name_testcase).bold = True

            # CT Summary.
            control = Main.wordSeachText(document=document, text=tag_language[3][otherConfigs['Language']])
            control.add_run(summary).bold = True

            # Executed by.
            control = Main.wordSeachText(document=document, text=tag_language[4][otherConfigs['Language']])
            control.add_run(full_name_run_test).bold = True

            # Evidence generate by.
            control = Main.wordSeachText(document=document, text=tag_language[5][otherConfigs['Language']])
            control.add_run(full_name_run_evidence).bold = True

            # Execution date
            control = Main.wordSeachText(document=document, text=tag_language[6][otherConfigs['Language']])
            control.add_run(str(completed_date)).bold = True

            # Total steps.
            control = Main.wordSeachText(document=document, text=tag_language[7][otherConfigs['Language']])
            control.add_run(str(step_number)).bold = True

            infoadd = True

        except Exception as ex:
            infoadd = False

            print(f"{Textcolor.FAIL}{logs['ErrorWordAddInfo']['Msg']}{Textcolor.END}", ex)
            Main.addLogs(message="General", value=logs["ErrorWordAddInfo"], value1=ex)

        return infoadd

    # Function to convert docx to pdf.
    def wordToPDF(**kwargs):

        try:
            # kwargs variables.
            path = kwargs.get("path")

            wdformatpdf = 17

            # Initialize.
            pythoncom.CoInitialize()

            word = win.Dispatch('Word.Application')
            document = word.Documents.Open(path)

            pdf_path = path.replace("docx", "pdf")
            document.SaveAs(pdf_path, FileFormat=wdformatpdf)
            document.Close()

            word.Quit()

            createpdf = pdf_path

        except Exception as ex:
            print(f"{Textcolor.FAIL}{logs['ErrorWordToPDF']['Msg']}{Textcolor.END}", ex)
            Main.addLogs(message="General", value=logs["ErrorWordToPDF"], value1=ex)
            createpdf = None

        return createpdf

    # Delete the files.
    def deleteFiles(**kwargs):

        try:
            # kwargs arguments.
            file_path = kwargs.get('path_log')
            extension = kwargs.get('extension')
            exact_file = kwargs.get('exact_file')

            # Delete all the files in a directory with the especific extension OR all if the extension is '*'.
            if file_path:

                files = os.listdir(file_path)

                for item in files:
                    if item.endswith(extension) or extension == '*':
                        os.remove(os.path.join(file_path, item))

            else:
                os.remove(exact_file)

        except Exception as ex:
            print(f"{Textcolor.FAIL}{logs['ErrorDeleteFiles']['Msg']}{Textcolor.END}", ex)
            Main.addLogs(message="General", value=logs["ErrorDeleteFiles"], value1=ex)

    # Add the log in the file.
    def addLogs(**kwargs):

        try:
            # kwargs variables.
            message = kwargs.get("message")
            value = kwargs.get("value")
            value1 = kwargs.get("value1")
            value2 = kwargs.get("value2")

            datetime_log: str = datetime.datetime.now().strftime("%d/%m/%y %H:%M:%S")

            # Get the hostname.
            otherConfigs["ComputerName"] = socket.gethostname()
            hostname = otherConfigs["ComputerName"]
            # Get the datetime.
            date_log = str(datetime.datetime.now().strftime("%d.%m.%Y"))
            # Set the file name.
            path = os.path.join(directories["LogFolder"], hostname + " - " + date_log + ".log")

            if not os.path.isdir(directories["LogFolder"]):
                Main.createDirectory(path_folder=directories["LogFolder"])

            # Append the log file.
            with open(path, 'a+', encoding='utf-8') as log_file:
                if message == 'NewConfig':  # Set the first line.
                    log_file.write("\n " + "*" * 61 + datetime_log + "*" * 61 + "\n")
                    log_file.write("\nLOAD INFORMATIONS FROM AZURE\n")
                elif message == 'NewSession':
                    log_file.write(str(value) + "\n")  # Test case name.
                    log_file.write("")
                elif message == 'General':
                    type_log = value['Type']
                    message_log = value['Msg']
                    local_log = value['Where']

                    # Alignment.
                    type_log = "{:<9}".format(type_log)

                    # Change the error origin if informed.
                    if value1 is not None and value2 is None:
                        msg_format = "{:<15}"
                        log_file.write(datetime_log + " - " + type_log + " - " + msg_format.format(message_log) + " - " + value1 + "\n")
                    elif value1 is not None and value2 is not None:
                        log_file.write(datetime_log + " - " + type_log + " - '" + value1 + "' " + message_log + " '" +
                                       value2 + "' - " + local_log + "\n")
                    else:
                        log_file.write(datetime_log + " - " + type_log + " - " + message_log + " - " + local_log + "\n")

                elif message == 'EndExecution':
                    log_file.write("\n " + "*" * 138 + "\n")
                else:
                    log_file.write(datetime_log + " - Log       - " + value + " " + " - " + str(value1) + "\n")

        except Exception as ex:
            Main.addLogs(message="General", value=logs["ErrorAddlog"]['Msg'], value1=str(ex))
            print(f"{Textcolor.FAIL}{logs['ErrorAddlog']['Msg']}{Textcolor.END}", ex)

    # Remove the HTML from the string.
    def removeHTML(**kwargs):

        try:
            # kwargs variables.
            value = kwargs.get("value")

            pattern = regex.compile('<.*?>')
            value = regex.sub(pattern, '', value).strip()

            value = value.replace('&nbsp;', ' ')
            value = value.replace('&lt;', '<')
            value = value.replace('&gt;', '>')

            return value

        except Exception as ex:

            print(f"{Textcolor.FAIL}{logs['ErrorRemoveHTML']['Msg']}{Textcolor.END}", ex)
            Main.addLogs(message="General", value=logs["ErrorRemoveHTML"]['Msg'], value1=str(ex))

    # Translate the messages.
    def translateMsg(**kwargs):
        try:
            # kwargs arguments.
            language = kwargs.get('language')
            new_hash = kwargs.get('new_hash')

            # Paths
            path_Origin_Yml = os.path.join(os.getcwd(), 'Automation', 'configs', 'dictionary-pt.yml')
            path_Translated_Yml = os.path.join(os.getcwd(), 'Automation', 'configs', 'dictionary-' + language + '.yml')

            regex_pattern = regex.compile('(.*?,)(.*?,)(.*)')

            tag = 'Msg:'

            if os.path.isfile(path_Translated_Yml):
                Main.deleteFiles(exact_file=path_Translated_Yml)

            with open(path_Origin_Yml, 'r') as yml_file:
                lines = yml_file.readlines()
                for line in lines:
                    if tag in line:
                        with open(path_Translated_Yml, 'a', encoding='utf-8') as yml_new_file:
                            groupType = regex.match(regex_pattern, line).group(1)
                            groupMsg = regex.match(regex_pattern, line).group(2)
                            msgTranslated = GoogleTranslator(source='pt', target=language).\
                                translate((groupMsg[len(tag) + 1:len(groupMsg) - 1]).strip())
                            groupWhere = regex.match(regex_pattern, line).group(3)
                            yml_new_file.write(f'{groupType}{tag} {msgTranslated},{groupWhere}\n')
                    else:
                        with open(path_Translated_Yml, 'a') as yml_new_file:
                            yml_new_file.write(f'{line}')

            Main.saveHash(new_hash=new_hash, path_part=language)
            print(f"{Textcolor.GREEN}{otherConfigs['TranslateMessage']}{Textcolor.END}")

        except Exception as ex:
            print(f"{Textcolor.FAIL}{logs['ErrorTranslateMessage']['Msg']}{Textcolor.END}", ex)
            Main.addLogs(message="General", value=logs["ErrorTranslateMessage"], value1=ex)

    # Configure the language for the automation.
    def configureLanguage(**kwargs):
        try:
            need_translation = False
            language = kwargs.get('language')

            path_file = os.path.join(os.getcwd(), directories["ConfigFolder"], 'dictionary-pt.yml')
            new_hash = Main.generateHash(path_file=path_file)
            actual_hash = Main.readHash(directory=directories["HashFolder"], language=language,
                                        actual_file='hash_dictionary.txt')

            if new_hash != actual_hash:
                need_translation = True

            return need_translation, new_hash

        except Exception as ex:
            print(f"{Textcolor.FAIL}{logs['ErrorConfigureLanguage']['Msg']}{Textcolor.END}", ex)
            Main.addLogs(message="General", value=logs["ErrorConfigureLanguage"], value1=ex)

    # Generate the hash for a file.
    def generateHash(**kwargs):
        try:
            # kwargs arguments.
            path_file = kwargs.get('path_file')

            BLOCKSIZE = 65536
            hasher = hashlib.sha1()
            with open(path_file, 'rb') as target_file:
                buffer = target_file.read(BLOCKSIZE)
                while len(buffer) > 0:
                    hasher.update(buffer)
                    buffer = target_file.read(BLOCKSIZE)

            return hasher.hexdigest()

        except Exception as ex:
            print(f"{Textcolor.FAIL}{logs['ErrorGenerateHash']['Msg']}{Textcolor.END}", ex)
            Main.addLogs(message="General", value=logs["ErrorGenerateHash"], value1=ex)

    # Read the hash in a file.
    def readHash(**kwargs):
        try:
            # kwargs arguments.
            directory = kwargs.get('directory')
            language = kwargs.get('language')
            actual_file = kwargs.get('actual_file')

            file_read = os.path.join(directory, language + '-' + actual_file)

            if os.path.exists(file_read):

                with open(file_read, 'r') as file:
                    content = file.readline()
            else:
                content = None

            return content

        except Exception as ex:
            print(f"{Textcolor.FAIL}{logs['ErrorReadHash']['Msg']}{Textcolor.END}", ex)
            Main.addLogs(message="General", value=logs["ErrorReadHash"], value1=ex)

    # Save the hash in a file.
    def saveHash(**kwargs):
        try:
            # kwargs arguments.
            new_hash = kwargs.get('new_hash')
            path_part = kwargs.get('path_part')

            # Variables.
            hash_file_path = os.path.join(directories["HashFolder"], path_part + '-hash_dictionary.txt')

            if not os.path.exists(directories['HashFolder']):
                Main.createDirectory(path_folder=directories['HashFolder'])
            with open(hash_file_path, 'w') as hash_file:
                hash_file.write(new_hash)

        except Exception as ex:
            print(f"{Textcolor.FAIL}{logs['ErrorSaveHash']['Msg']}{Textcolor.END}", ex)
            Main.addLogs(message="General", value=logs["ErrorSaveHash"], value1=ex)

    # Load the configuration file.
    def loadConfigs(**kwargs):

        # kwargs parameters.
        language = kwargs.get('language')

        try:
            global verbs, logs, directories, otherConfigs, searchForAttribute, searchForComponent

            # yml path.
            path = os.path.join(os.getcwd(), 'Automation', 'configs', 'dictionary-' + language + '.yml')
            if not os.path.exists(path):  # For the exe file.
                path = os.path.join(os.getcwd(), 'configs', 'dictionary-' + language + '.yml')

            with open(path, encoding='utf-8') as configFile:
                config = yaml.safe_load(configFile)

            # Add the sections.
            verbs = config['verbs']
            logs = config['logs']
            directories = config['directories']
            otherConfigs = config['otherConfigs']
            searchForAttribute = config["searchForAttribute"]
            searchForComponent = config["searchForComponent"]

        except Exception as ex:
            print(f"{Textcolor.FAIL}{logs['ErrorLoadConfigs']['Msg']}{Textcolor.END}", ex)
            Main.addLogs(message="General", value=logs["ErrorLoadConfigs"], value1=str(ex))

    # Calculate the percentage of execution.
    def percentage(**kwargs):
        try:
            # kwargs parameters.
            actual_number = kwargs.get('actual')
            total = kwargs.get('total')

            line = "*" * 12
            percentage = "{0:.2f}".format((actual_number / total) * 100)

            print(f"{Textcolor.BLUE}{' '}"
                  f"{otherConfigs['Percentage']['Msg']}"
                  f"{' '}{Textcolor.UNDERLINE}{percentage}{'%'}{' '}{Textcolor.END}"
                  f"{Textcolor.BLUE}{Textcolor.END}\n")

            Main.addLogs(message="General", value=logs["Percentage"])

        except Exception as ex:
            print(f"{Textcolor.FAIL}{logs['ErrorPercentage']['Msg']}{Textcolor.END}", ex)
            Main.addLogs(message="General", value=logs["ErrorPercentage"], value1=str(ex))

    # Verify if the file exist.
    def verifyFile(**kwargs):
        try:

            # kwargs parameters.
            path = kwargs.get('path')
            extension = kwargs.get('extension')
            msg_not_found = kwargs.get('msg_not_found')
            msg_found = kwargs.get('msg_found')

            files = os.listdir(path)
            for file in files:
                if file.endswith(extension):
                    print(f"{Textcolor.GREEN}{msg_found}{Textcolor.END}")
                    return True
                else:
                    print(f"{Textcolor.GREEN}{msg_not_found}{Textcolor.END}")
                    return False

        except Exception as ex:
            print(f"{Textcolor.FAIL}{logs['ErrorVerifyFile']['Msg']}{Textcolor.END}", ex)
            Main.addLogs(message="General", value=logs["ErrorVerifyFile"], value1=str(ex))

    # Compare files using Beyond Compare.
    def compareBeyondCompare(**kwargs):
        try:
            # kwargs parameters.
            baseline = kwargs.get('baseline')
            new_file = kwargs.get('new_file')
            test_name = kwargs.get('test_name')

            baseline = os.path.join(directories['CompareDownloadFolder'], test_name, baseline)
            new_file = os.path.join(directories['CompareDownloadFolder'], test_name, new_file)

            # Update the Settings.
            subprocess.Popen([directories['BeyondCompare'], directories['BeyondCompareSettings'], '/silent'])

            file_session = Main._checkSessionBCFile(baseline=baseline)

            # Open the session.
            subprocess.Popen([directories['BeyondCompare'], file_session, '/silent'])

            # Disable the mouse.
            pyautogui.FAILSAFE = False

            # Open the file (Left).
            time.sleep(10)
            pyautogui.hotkey('ctrl', 'o', interval=.5)
            pyautogui.typewrite(baseline)
            pyautogui.hotkey('Enter', interval=.5)

            pyautogui.press('tab', interval=.5)

            # Open the file (Right).
            pyautogui.hotkey('ctrl', 'o', interval=.5)
            pyautogui.typewrite(new_file)
            pyautogui.hotkey('Enter', interval=.5)

            # Enable the edit option.
            pyautogui.press('f2', interval=.5)

            # Disable the mouse.
            pyautogui.FAILSAFE = True

            while Main.checkProcess(process='BCompare.exe'):
                time.sleep(1)

            print(f"{Textcolor.GREEN}{test_name} - {logs['CompareFile']['Msg']}{Textcolor.END}")

        except Exception as ex:
            print(f"{Textcolor.FAIL}{logs['ErrorCompareFile']['Msg']}{Textcolor.END}", ex)
            Main.addLogs(message="General", value=logs["ErrorCompareFile"], value1=str(ex))

    # Check if the process is running.
    def checkProcess(**kwargs):

        # kwargs variables.
        process = kwargs.get("process")

        # List the running process.
        output = os.popen('wmic process get description, processid').read()

        if process in output:
            return True
        else:
            return False

    # Load the session file.
    def _checkSessionBCFile(**kwargs):

        # kwargs variable.
        file = kwargs.get("baseline")

        if file.upper().endswith('CSV'):
            return 'CSVFiles'
        elif file.upper().endswith('TXT'):
            return 'TXTFiles'
        elif file.upper().endswith('PDF'):
            return 'PDFFiles'
        elif file.upper().endswith('DOC') or file.upper().endswith('DOCX'):
            return 'WordFiles'
        elif file.upper().endswith('GIF') or file.upper().endswith('ICO') or file.upper().endswith('JPG') or \
                file.upper().endswith('PNG') or file.upper().endswith('TIF') or file.upper().endswith('BMP'):
            return 'ImageFiles'
        elif file.upper().endswith('MP4'):
            return 'VideoFiles'

    # Check if the test case is a Desktop test case.
    def _checkDesktop_TC(**kwargs):

        # kwargs variable.
        list_steps = kwargs.get("list_steps")

        desktop_TC = False

        for cont, _ in enumerate(otherConfigs['DesktopFunctions']):
            if any(otherConfigs['DesktopFunctions'][cont] in desktop_verb for desktop_verb in list_steps):
                desktop_TC = True

        return desktop_TC

    # Check for Updates.
    # def check_Updates(self,):
    #
    #     try:
    #         self.URL = 'https://' + otherConfigs['Token_GitHub'] + otherConfigs['GitHubReadMe']
    #
    #         headers = {
    #             'Authorization': f"token {otherConfigs['Token_GitHub']}",
    #             'Accept': 'application/vnd.github.v4+raw'
    #         }
    #
    #         if os.path.exists(directories['ReadMeFile']):
    #             with open(directories['ReadMeFile'], 'r') as readme:
    #                 Main.addLogs(None, message="General", value=logs["UpdateNoNewVersion"], value1="ReadMe") ### Só para testar, pode retirar
    #
    #                 for line in readme:
    #                     if 'Version' in line:
    #                         LocalVersion = regex.search('Version(.*)\*\*', line).group(1)
    #                         LocalVersion = LocalVersion.strip()
    #                         break
    #         else:
    #             with open(directories['ReadMeFileEXEC'], 'r') as readme:
    #                 Main.addLogs(None, message="General", value=logs["UpdateNoNewVersion"], value1="ReadMeEXEC") ### Só para testar, pode retirar
    #
    #                 for line in readme:
    #                     if 'Version' in line:
    #                         LocalVersion = regex.search('Version(.*)\*\*', line).group(1)
    #                         LocalVersion = LocalVersion.strip()
    #                         break
    #
    #         response = requests.get(self.URL, headers=headers)
    #
    #         if response.status_code == 200:
    #             GitHubVersion = response.text
    #             GitHubVersion = regex.search('Version(.*)\*\*', response.text).group(1)
    #             GitHubVersion = GitHubVersion.strip()
    #
    #             # If the GitHub version is latest than local version.
    #             if GitHubVersion > LocalVersion:
    #                 # update.show_Update_input() # CORRIGIR
    #                 # update.run() # CORRIGIR
    #                 pass
    #
    #             else:
    #                 print(f"{Textcolor.GREEN}{logs['UpdateNoNewVersion']['Msg']}{Textcolor.END}")
    #                 Main.addLogs(message="General", value=logs["UpdateNoNewVersion"],
    #                              value1=logs["UpdateNoNewVersion"]["Msg"])
    #
    #         else:
    #             print(f"{Textcolor.FAIL}{logs['CouldNotCheckForUpdates']['Msg']}{Textcolor.END}")
    #             Main.addLogs(message="General", value=logs["CouldNotCheckForUpdates"],
    #                          value1=logs["CouldNotCheckForUpdates"]["Msg"])
    #             #update.could_not_check_for_updates_msg() # CORRIGIR
    #             #update.run() # CORRIGIR
    #
    #     except Exception as ex:
    #         print(f"{Textcolor.FAIL}{logs['DownloadUpdateFunctionFailed']['Msg']}{Textcolor.END}", ex)
    #         Main.addLogs(message="General", value=logs["ErrorStartAutomation"], value1=str(ex))
    #
    # # Check for new automation version.
    # def download_Updates(self):
    #     try:
    #         headers = {
    #             'Authorization': f"token {otherConfigs['Token_GitHub']}",
    #             'Accept': 'application/vnd.github.v4+raw'
    #         }
    #
    #         # Download folder.
    #         DOWNLOAD_DIR = directories['UpdateFolder']
    #         LINK_DOWNLOAD = 'https://' + otherConfigs['Token_GitHub'] + otherConfigs['GitHubContent']
    #         local_filename = LINK_DOWNLOAD.split('/')[-1]
    #
    #         # Execute the new version download.
    #         with requests.get(LINK_DOWNLOAD, stream=True, headers=headers) as r:
    #             if r.status_code == 200:
    #
    #                 # Create the New Version directory.
    #                 Main.createDirectory(path_folder=directories['UpdateFolder'])
    #
    #                 with open(DOWNLOAD_DIR + '\\' + local_filename, 'wb') as f:
    #                     shutil.copyfileobj(r.raw, f)
    #                     print(f"{Textcolor.GREEN}{logs['DownloadPackageCompleted']['Msg']}{Textcolor.END}")
    #                     Main.addLogs(message="General", value=logs["DownloadPackageCompleted"],
    #                                  value1=logs["DownloadPackageCompleted"]["Msg"])
    #
    #             else:
    #                 print(f"{Textcolor.FAIL}{logs['ErrorDownloadUpdate']['Msg']}{Textcolor.END}", value1=r.status_code)
    #                 # update.download_update_fail_msg() # CORRIGIR
    #                 # update.run() # CORRIGIR
    #
    #         LINK_DOWNLOAD_BAT = 'https://' + otherConfigs['Token_GitHub'] + otherConfigs['GitHubBatFile']
    #         local_filename = LINK_DOWNLOAD_BAT.split('/')[-1]
    #
    #         # Download the BAT file.
    #         with requests.get(LINK_DOWNLOAD_BAT, stream=True, headers=headers) as r:
    #             if r.status_code == 200:
    #
    #                 with open(DOWNLOAD_DIR + '\\' + local_filename, 'wb') as f:
    #                     shutil.copyfileobj(r.raw, f)
    #                     print(f"{Textcolor.GREEN}{logs['DownloadBATCompleted']['Msg']}{Textcolor.END}")
    #                     Main.addLogs(message="General", value=logs["DownloadBATCompleted"],
    #                                  value1=logs["DownloadBATCompleted"]["Msg"])
    #             else:
    #                 print(f"{Textcolor.FAIL}{logs['ErrorDownloadUpdate']['Msg']}{Textcolor.END}")
    #                 # update.download_update_fail_msg() # CORRIGIR
    #                 # update.run() # CORRIGIR
    #
    #         # update.download_update_completed_msg.run() # CORRIGIR
    #
    #     except Exception as ex:
    #         print(f"{Textcolor.FAIL}{logs['DownloadUpdateFunctionFailed']['Msg']}{Textcolor.END}", ex)
    #         Main.addLogs(message="General", value=logs["ErrorStartAutomation"], value1=str(ex))
    #
    # def install_Update(self):
    #     try:
    #         subprocess.Popen([directories['UpdateFolder'], otherConfigs["InstallBAT"]])
    #
    #         shutil.rmtree(os.path.join(directories['UpdateFolder']))
    #
    #         print(f"{Textcolor.GREEN}{logs['InstallNewVersion']['Msg']}{Textcolor.END}")
    #         Main.addLogs(message="General", value=logs["InstallNewVersion"], value1=logs["InstallNewVersion"]["Msg"])
    #
    #     except Exception as ex:
    #         print(f"{Textcolor.FAIL}{logs['ErrorInstallNewVersion']['Msg']}{Textcolor.END}", ex)
    #         Main.addLogs(message="General", value=logs["ErrorInstallNewVersion"], value1=str(ex))
